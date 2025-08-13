#!/usr/bin/env python3
"""
Create ASPCA test document in .docx format
"""

try:
    from docx import Document
    from docx.shared import Inches
    
    def create_aspca_docx():
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('ASPCA Test Document', 0)
        
        # Add content
        doc.add_heading('About ASPCA', level=1)
        
        p1 = doc.add_paragraph(
            'The American Society for the Prevention of Cruelty to Animals (ASPCA) is a '
            'non-profit organization dedicated to preventing animal cruelty. Founded in 1866, '
            'the ASPCA was the first animal welfare organization established in North America.'
        )
        
        doc.add_heading('Mission Statement', level=2)
        mission = doc.add_paragraph(
            'To provide effective means for the prevention of cruelty to animals throughout the United States.'
        )
        
        doc.add_heading('Services', level=2)
        services = doc.add_paragraph()
        services.add_run('• Animal rescue and placement\n')
        services.add_run('• Animal health services\n')
        services.add_run('• Anti-cruelty investigations\n')
        services.add_run('• Government relations and advocacy\n')
        services.add_run('• Animal behavior and training')
        
        doc.add_heading('Contact Information', level=2)
        contact = doc.add_paragraph()
        contact.add_run('Website: ').bold = True
        contact.add_run('https://www.aspca.org\n')
        contact.add_run('Phone: ').bold = True
        contact.add_run('(212) 876-7700\n')
        contact.add_run('Address: ').bold = True
        contact.add_run('424 E 92nd St, New York, NY 10128-6804')
        
        doc.add_heading('Key Programs', level=2)
        programs = doc.add_paragraph()
        programs.add_run('1. ASPCA Animal Poison Control Center\n')
        programs.add_run('2. ASPCA Behavioral Rehabilitation Center\n')
        programs.add_run('3. ASPCA Spay/Neuter Alliance\n')
        programs.add_run('4. ASPCA Field Investigation and Response team')
        
        conclusion = doc.add_paragraph(
            'The ASPCA operates the nation\'s poison control center for animals and provides '
            '24/7 emergency assistance. They also run adoption centers and provide educational '
            'resources for pet owners. The organization is committed to advancing the safety '
            'and well-being of animals through advocacy, education, and hands-on care.'
        )
        
        # Save document
        doc.save('ASPCATest.docx')
        print("✅ Created ASPCATest.docx successfully")
        return True
        
    if __name__ == "__main__":
        create_aspca_docx()
        
except ImportError:
    print("⚠️  python-docx not installed. Will create after installing requirements.")
    print("Run: pip install python-docx")